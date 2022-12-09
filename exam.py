import difflib
import io
import json
import random

import docx
from docx.document import Document
from docx.oxml.ns import qn

META_DATA = {
    'list': [
        {
            'id': 'sixiu',
            'name': '思修',
            'pages': [
                'sixiu_0',
                'sixiu_1',
                'sixiu_2',
                'sixiu_3',
                'sixiu_4',
                'sixiu_5',
                'sixiu_6'
            ]
        },
        {
            'id': 'jindaishi',
            'name': '近代史',
            'pages': [
                'jindaishi_0',
                'jindaishi_1',
                'jindaishi_2',
                'jindaishi_3',
                'jindaishi_4',
                'jindaishi_5',
                'jindaishi_6',
                'jindaishi_7',
                'jindaishi_8',
                'jindaishi_9',
                'jindaishi_10',
                'jindaishi_11',
                'jindaishi_12',
                'jindaishi_13'
            ]
        },
        {
            'id': 'mayuan',
            'name': '马原',
            'pages': [
                'mayuan_0',
                'mayuan_1',
                'mayuan_2',
                'mayuan_3',
                'mayuan_4',
                'mayuan_5',
                'mayuan_6',
                'mayuan_7'
            ]
        },
        {
            'id': 'maogai',
            'name': '毛概',
            'pages': [
                'maogai_0',
                'maogai_1',
                'maogai_2',
                'maogai_3',
                'maogai_4',
                'maogai_5',
                'maogai_6',
                'maogai_7',
                'maogai_8',
                'maogai_9',
                'maogai_10',
                'maogai_11',
                'maogai_12',
                'maogai_13'
            ]
        }
    ],
    'map': {
        'sixiu_0': '导论',
        'sixiu_1': '第一章',
        'sixiu_2': '第二章',
        'sixiu_3': '第三章',
        'sixiu_4': '第四章',
        'sixiu_5': '第五章',
        'sixiu_6': '第六章',
        'jindaishi_0': '上篇综述',
        'jindaishi_1': '第一章',
        'jindaishi_2': '第二章',
        'jindaishi_3': '第三章',
        'jindaishi_4': '中篇综述',
        'jindaishi_5': '第四章',
        'jindaishi_6': '第五章',
        'jindaishi_7': '第六章',
        'jindaishi_8': '第七章',
        'jindaishi_9': '下篇综述',
        'jindaishi_10': '第八章',
        'jindaishi_11': '第九章',
        'jindaishi_12': '第十章',
        'jindaishi_13': '第十一章',
        'mayuan_0': '导论',
        'mayuan_1': '第一章',
        'mayuan_2': '第二章',
        'mayuan_3': '第三章',
        'mayuan_4': '第四章',
        'mayuan_5': '第五章',
        'mayuan_6': '第六章',
        'mayuan_7': '第七章',
        'maogai_0': '第一章',
        'maogai_1': '第二章',
        'maogai_2': '第三章',
        'maogai_3': '第四章',
        'maogai_4': '第五章',
        'maogai_5': '第六章',
        'maogai_6': '第七章',
        'maogai_7': '第八章',
        'maogai_8': '第九章',
        'maogai_9': '第十章',
        'maogai_10': '第十一章',
        'maogai_11': '第十二章',
        'maogai_12': '第十三章',
        'maogai_13': '第十四章'
    }
}

# 需要保留的数据
PROP_FILTER = {
    'questionId',
    'courseId',
    'questionStem',
    'questionType',
    'questionTypeName',
    'courseLineId',
    'courseLineNumber',
    'answer',
    'options'
}

# 题目数据
QUES_DATA = {'index': {}}
# 题目类型
QUES_TYPE = ['radio', 'checkbox', 'fillblank', 'judgment']
# 生成试卷题目占比
QUES_EXAM = [60, 0, 20, 20]

for course in META_DATA['list']:
    for line in course['pages']:
        with open('./data/{}.json'.format(line), encoding='utf-8') as f:
            file_content = json.load(f)
        QUES_DATA[line] = []
        for q in file_content['body']['list']:
            if q['questionType'] == 'checkbox':
                continue
            QUES_DATA[line].append(q)
            for key in q.keys() - PROP_FILTER:
                q.pop(key)
            QUES_DATA['index'][q['questionId']] = q
    course_name = course['id']
    for t in QUES_TYPE:
        QUES_DATA[course_name + '_' + t] = []
    for line in course['pages']:
        for q in QUES_DATA[line]:
            if q['questionType'] in ('radio', 'checkbox', 'fillblank'):
                QUES_DATA[course_name + '_' + q['questionType']].append(q)
            else:
                QUES_DATA[course_name + '_judgment'].append(q)


class Exam:
    ''' 学校思政云系统模拟测试 '''
    @staticmethod
    def get_que(line: str, num: int) -> dict:
        ''' 获取题目 '''
        return QUES_DATA[line][num]

    @staticmethod
    def get_line(line: str) -> list:
        ''' 获取章节 '''
        return QUES_DATA[line]

    @staticmethod
    def make_exam(course: str, seed: int):
        ''' 生成考试试卷 '''
        rnd = random.Random(None)
        rnd.seed(seed)
        res = []
        for t, num in zip(QUES_TYPE, QUES_EXAM):
            res.extend(rnd.sample(QUES_DATA[course + '_' + t], num))
        return res

    @staticmethod
    def query_que(text: str) -> dict:
        ''' 搜题 '''
        curr = -0.1
        res = {}
        for line in META_DATA['map']:
            for q in QUES_DATA[line]:
                diff = difflib.SequenceMatcher(
                    None, text, q['questionStem']).quick_ratio()
                if diff > curr:
                    curr = diff
                    res = q
        return res

    @staticmethod
    def search_for(ids: list) -> list:
        ''' 返回列表内的题目 '''
        return [QUES_DATA['index'][qid] for qid in ids]

    @staticmethod
    def make_docx(ids: list, title: str):
        ''' 生成 docx 文件 '''
        doc: Document = docx.Document()
        doc.add_heading(title, 0)
        doc.styles['Normal'].font.name = 'Cambria'
        doc.styles['Normal']._element.rPr.rFonts.set(
            qn('w:eastAsia'), 'SimSun-ExtB')
        filestream = io.BytesIO()
        for i, q in enumerate(Exam.search_for(ids)):
            doc.add_paragraph('{}. {}'.format(i + 1, q['questionStem']))
            if q['questionType'] in ('radio', 'checkbox'):
                for opt in q['options']:
                    doc.add_paragraph(opt['option'] + '. ' +
                                      opt['content'].strip())
            elif q['questionType'] == 'fillblank':
                for opt in q['options']:
                    doc.add_paragraph(
                        opt['option'] + ': ' + opt['content'].strip())
            if q['questionType'] != 'fillblank':
                doc.add_paragraph('答案：' + q['answer'])
        doc.save(filestream)
        return filestream


if __name__ == '__main__':
    for course in META_DATA['list']:
        course_name = course['name']
        for line in course['pages']:
            line_name = META_DATA['map'][line]
            file_name = '{} - {}'.format(course_name, line_name)
            qid = [q['questionId'] for q in Exam.get_line(line)]
            bytes_io = Exam.make_docx(qid, file_name)
            with open(file_name + '.docx', 'wb') as f:
                f.write(bytes_io.getvalue())
            print('{}.docx 已生成'.format(file_name))
